#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Автоматическое заполнение таблиц аудита - УЛУЧШЕННАЯ ВЕРСИЯ v2.3
Использует Google Gemini Vision API для обработки документов

УЛУЧШЕНИЯ:
✅ Полное распознавание текста (не обрывается)
✅ Умное сопоставление вопросов-ответов
✅ Постобработка OCR для исправления ошибок
✅ Улучшенные промпты для AI
✅ Максимальный контекст для AI (500 символов на вопрос, до 100 вопросов)
✅ Фильтрация OCR-артефактов (лишние символы)
✅ Классификация документов и извлечение только нужных фрагментов
✅ Столбец D с именем файла-источника
✅ Поддержка DOC/DOCX/PDF (не только изображения)
"""

import tkinter as tk
from tkinter import ttk, filedialog, messagebox, scrolledtext
import os
import sys
from pathlib import Path
import threading
import json
from datetime import datetime
import subprocess
import platform
import re

# Проверка и импорт необходимых библиотек
try:
    import requests
except ImportError:
    print("❌ Ошибка: модуль 'requests' не установлен")
    print("Установите: pip install requests")
    sys.exit(1)

try:
    from openpyxl import Workbook, load_workbook
except ImportError:
    print("❌ Ошибка: модуль 'openpyxl' не установлен")
    print("Установите: pip install openpyxl")
    sys.exit(1)


class TextPostProcessor:
    """Постобработчик для исправления типичных ошибок OCR"""

    @staticmethod
    def fix_ocr_errors(text):
        """Исправляет типичные ошибки OCR для русского языка и удаляет артефакты"""
        if not text:
            return text

        # УДАЛЕНИЕ OCR-АРТЕФАКТОВ
        # URL-кодировки
        text = re.sub(r'%[0-9A-Fa-f]{2}', '', text)

        # Типичные мусорные последовательности
        artifacts = [
            'EMV NOT', 'POLATOM', 'РОСАТОМ POLATOM',
            'НРБ 10-26', 'ДБ EMV', 'Курсовой 10-31',
            '6-10 НРБ', 'жденные%', 'npc курсатов',
            '1 / 2 90%', 'п/п |', '| :----',
        ]
        for artifact in artifacts:
            text = text.replace(artifact, '')

        # Удаление путей файлов (содержащих / или \)
        text = re.sub(r'[a-zA-Zа-яА-Я0-9_\-]+[/\\][a-zA-Zа-яА-Я0-9_\-/\\%]+', '', text)

        # Удаление повторяющихся символов (например, "----", "====")
        text = re.sub(r'([=\-|_])\1{5,}', '', text)

        # Исправляем разделенные буквы (О О О → ООО)
        text = re.sub(r'О\s+О\s+О', 'ООО', text)
        text = re.sub(r'П\s+Р\s+И\s+К\s+А\s+З', 'ПРИКАЗ', text)
        text = re.sub(r'У\s+Д\s+О\s+С\s+Т\s+О\s+В\s+Е\s+Р\s+Е\s+Н\s+И\s+Е', 'УДОСТОВЕРЕНИЕ', text)

        # Исправляем даты
        text = re.sub(r'(\d{2})\s*\.\s*(\d{2})\s*\.\s*(\d{4})', r'\1.\2.\3', text)

        # Убираем лишние пробелы
        text = re.sub(r'\s+', ' ', text)
        text = text.strip()

        # Восстанавливаем параграфы
        text = re.sub(r'\.\s+([А-ЯA-Z])', r'.\n\n\1', text)

        return text

    @staticmethod
    def extract_metadata(text):
        """Извлекает метаданные из текста (организация, дата, номер и т.д.)"""
        metadata = {
            'organization': None,
            'doc_type': None,
            'doc_number': None,
            'doc_date': None,
            'persons': []
        }

        # Организация
        org_match = re.search(r'(?:Общество с ограниченной ответственностью|ООО)\s*[«"]?([^»"]+)[»"]?', text, re.IGNORECASE)
        if org_match:
            metadata['organization'] = org_match.group(1).strip()

        # Тип документа
        for doc_type in ['ПРИКАЗ', 'УДОСТОВЕРЕНИЕ', 'СПРАВКА', 'АКТ', 'ПРОТОКОЛ']:
            if doc_type in text.upper():
                metadata['doc_type'] = doc_type
                break

        # Номер документа
        num_match = re.search(r'№\s*(\d+[\d\/\-]*)', text)
        if num_match:
            metadata['doc_number'] = num_match.group(1)

        # Дата
        date_match = re.search(r'(\d{2}\.\d{2}\.\d{4})', text)
        if date_match:
            metadata['doc_date'] = date_match.group(1)

        # ФИО
        name_pattern = r'([А-ЯЁ][а-яё]+)\s+([А-ЯЁ][а-яё]+)\s+([А-ЯЁ][а-яё]+)'
        names = re.findall(name_pattern, text)
        metadata['persons'] = [' '.join(name) for name in names[:5]]  # Максимум 5 имен

        return metadata

    @staticmethod
    def classify_document_type(text):
        """
        Улучшенная классификация типа документа на основе содержимого

        Возвращает один из типов:
        - 'regulation' - приказ, положение, инструкция, процедура
        - 'schedule' - график
        - 'certificate' - документ об обучении (удостоверение)
        - 'unknown' - не удалось определить
        """
        text_upper = text.upper()

        # Подсчет веса для каждого типа документа
        scores = {
            'certificate': 0,
            'schedule': 0,
            'regulation': 0
        }

        # Документы об обучении - более точные критерии
        certificate_keywords = {
            'УДОСТОВЕРЕНИЕ': 3,
            'ПОВЫШЕНИИ КВАЛИФИКАЦИИ': 3,
            'ПОВЫШЕНИЕ КВАЛИФИКАЦИИ': 3,
            'ОБУЧЕНИЕ': 1,
            'ПРОГРАММЕ': 1,
            'ПРОШЕЛ ОБУЧЕНИЕ': 2,
            'ПРОШЛА ОБУЧЕНИЕ': 2,
            'ОБЪЕМЕ': 1,
            'ЧАСОВ': 1,
            'АКАДЕМИЯ': 2,
            'УЧЕБНЫЙ ЦЕНТР': 2,
            'КУРС': 1
        }

        # Графики - улучшенные критерии
        schedule_keywords = {
            'ГРАФИК': 3,
            'РАСПИСАНИЕ': 3,
            'ПЛАН-ГРАФИК': 4,
            'МЕРОПРИЯТИЕ': 1,
            'ОТВЕТСТВЕННОЕ ЛИЦО': 2,
            'СРОК ВЫПОЛНЕНИЯ': 2,
            'МЕСЯЦ': 1,
            'КВАРТАЛ': 1
        }

        # Нормативные документы - улучшенные критерии
        regulation_keywords = {
            'ПРИКАЗ': 3,
            'ПРИКАЗЫВАЮ': 4,
            'ПОЛОЖЕНИЕ': 2,
            'ИНСТРУКЦИЯ': 2,
            'ПРОЦЕДУРА': 2,
            'РЕГЛАМЕНТ': 2,
            'УТВЕРДИТЬ': 2,
            'НАЗНАЧИТЬ': 2,
            'ДОВЕСТИ ДО СВЕДЕНИЯ': 2,
            'ГЕНЕРАЛЬНЫЙ ДИРЕКТОР': 2,
            'РУКОВОДИТЕЛЬ': 1
        }

        # Подсчет баллов для каждого типа
        for keyword, weight in certificate_keywords.items():
            if keyword in text_upper:
                scores['certificate'] += weight

        for keyword, weight in schedule_keywords.items():
            if keyword in text_upper:
                scores['schedule'] += weight

        for keyword, weight in regulation_keywords.items():
            if keyword in text_upper:
                scores['regulation'] += weight

        # Определяем тип с наибольшим баллом
        max_score = max(scores.values())

        if max_score == 0:
            return 'unknown'

        # Возвращаем тип с максимальным баллом
        for doc_type, score in scores.items():
            if score == max_score:
                return doc_type

        return 'unknown'

    @staticmethod
    def smart_truncate(text, max_length=500):
        """
        Умное обрезание текста с сохранением целостности предложений

        Пытается обрезать на конце предложения (. ! ?) или на границе слова
        """
        if len(text) <= max_length:
            return text

        # Пытаемся найти конец предложения в пределах max_length
        truncated = text[:max_length]

        # Ищем последнюю точку, восклицательный или вопросительный знак
        for end_char in ['. ', '! ', '? ']:
            last_pos = truncated.rfind(end_char)
            if last_pos > max_length * 0.5:  # Хотя бы половина текста должна остаться
                return truncated[:last_pos + 1].strip()

        # Если не нашли конец предложения, обрезаем на границе слова
        last_space = truncated.rfind(' ')
        if last_space > 0:
            return truncated[:last_space].strip() + '...'

        # В крайнем случае просто обрезаем
        return truncated.strip() + '...'

    @staticmethod
    def extract_relevant_fragment(text, doc_type, metadata):
        """
        Извлекает только нужный фрагмент документа в соответствии с правилами:

        - Положения/инструкции/процедуры: название, шифр, дата
        - Графики: название, дата, столбцы
        - Документы об обучении: тема, даты, № удостоверения (БЕЗ ФИО!)
        """
        if doc_type == 'certificate':
            # Документ об обучении - БЕЗ ФИО!
            parts = []

            # Тема обучения (ищем после "программе:")
            theme_match = re.search(r'программе[:\s]+[«"]?([^»"\n]+?)[»"\n]', text, re.IGNORECASE)
            if theme_match:
                theme = TextPostProcessor.smart_truncate(theme_match.group(1).strip(), 300)
                parts.append(f"Тема: {theme}")

            # Даты обучения
            dates = re.findall(r'(\d{2}\.\d{2}\.\d{4})', text)
            if len(dates) >= 2:
                parts.append(f"Даты: с {dates[0]} по {dates[1]}")
            elif len(dates) == 1:
                parts.append(f"Дата: {dates[0]}")

            # Номер удостоверения (ищем после "№")
            cert_num_match = re.search(r'(?:удостоверение|№)\s*([A-ZА-Я0-9\-/]+)', text, re.IGNORECASE)
            if cert_num_match:
                parts.append(f"№: {cert_num_match.group(1)}")

            return '\n'.join(parts) if parts else text[:300]

        elif doc_type == 'regulation':
            # Приказ/Положение/Инструкция
            parts = []

            # Полное название
            if metadata.get('doc_type'):
                parts.append(metadata['doc_type'])

            # Шифр/номер
            if metadata.get('doc_number'):
                parts.append(f"№ {metadata['doc_number']}")

            # Дата
            if metadata.get('doc_date'):
                parts.append(f"от {metadata['doc_date']}")

            # Название (ищем после "О " или "Об " до точки или конца абзаца)
            title_match = re.search(r'(?:О|Об)\s+([^\n]+?)(?:\.|$)', text, re.IGNORECASE)
            if title_match:
                title = TextPostProcessor.smart_truncate(title_match.group(0).strip(), 400)
                parts.append(f"Название: {title}")

            return '\n'.join(parts) if parts else text[:300]

        elif doc_type == 'schedule':
            # План/График
            parts = []

            # Название плана/графика (ищем до точки или конца строки)
            title_match = re.search(r'(?:график|план)[:\s]+([^\n]+?)(?:\.|$)', text, re.IGNORECASE)
            if title_match:
                title = TextPostProcessor.smart_truncate(title_match.group(0).strip(), 400)
                parts.append(f"Название: {title}")

            # Дата утверждения
            if metadata.get('doc_date'):
                parts.append(f"Дата утверждения: {metadata['doc_date']}")

            # Перечень столбцов (ищем строки с "|" - таблица)
            table_lines = [line for line in text.split('\n') if '|' in line]
            if table_lines and len(table_lines) > 0:
                # Первая строка - заголовки столбцов
                parts.append(f"Столбцы: {table_lines[0][:200]}")

                # Ищем пример мероприятия (первая строка после заголовка с данными)
                for line in table_lines[1:]:
                    # Пропускаем разделители (строки типа |---|---|)
                    if not re.match(r'^\s*\|[\s\-:]+\|\s*$', line):
                        parts.append(f"Пример мероприятия: {line[:200]}")
                        break

            return '\n'.join(parts) if parts else text[:300]

        # Неизвестный тип - возвращаем первые 300 символов
        return text[:300]


class AuditProcessorApp:
    """Главное приложение для обработки аудита"""

    def __init__(self, root):
        self.root = root
        self.root.title("🔍 Audit Processor v2.3 - Автозаполнение таблиц аудита (УЛУЧШЕННАЯ ВЕРСИЯ)")
        self.root.geometry("900x700")
        self.root.configure(bg="#f5f5f5")

        # Загрузка конфигурации
        self.load_config()

        # Инициализация AI провайдера
        self.init_ai_provider()

        # Инициализация постобработчика
        self.post_processor = TextPostProcessor()

        self.setup_ui()

    def load_config(self):
        """Загрузка конфигурации из config.json"""
        config_path = Path(__file__).parent / "config.json"

        if config_path.exists():
            try:
                with open(config_path, 'r', encoding='utf-8') as f:
                    self.config = json.load(f)
                print(f"✅ Конфигурация загружена из {config_path}")
            except Exception as e:
                print(f"❌ Ошибка загрузки config.json: {e}")
                self.config = {"ai_provider": "gemini"}
        else:
            print("❌ config.json не найден! Создайте config.json с Google Gemini API ключом")
            self.config = {"ai_provider": "gemini"}

    def init_ai_provider(self):
        """Инициализация Google Gemini Vision API"""
        self.ai_provider = "gemini"

        # Google Gemini
        gemini_config = self.config.get("gemini", {})
        self.gemini_api_key = gemini_config.get("api_key")
        self.gemini_model = gemini_config.get("model", "gemini-1.5-flash")

        if self.gemini_api_key:
            try:
                import google.generativeai as genai
                genai.configure(api_key=self.gemini_api_key)
                model_name = self.gemini_model.replace("models/", "")
                self.gemini_client = genai.GenerativeModel(model_name)
                self.ai_available = True
                print(f"✅ Google Gemini подключен ({model_name})")
                print("🎉 Обработка будет быстрой и качественной!")
            except Exception as e:
                error_msg = str(e)
                print(f"❌ Ошибка подключения Gemini: {error_msg}")
                # Проверяем на ошибки лимитов API
                if "quota" in error_msg.lower() or "limit" in error_msg.lower():
                    print("⚠️ Возможно исчерпан лимит API Gemini. Проверьте квоту на https://aistudio.google.com")
                self.ai_available = False
        else:
            print("❌ API ключ Gemini не найден в config.json")
            print("💡 Добавьте в config.json:")
            print('   {"ai_provider": "gemini", "gemini": {"api_key": "YOUR_KEY", "model": "gemini-1.5-flash"}}')
            self.ai_available = False

    def setup_ui(self):
        """Создание интерфейса"""

        # Заголовок
        header = tk.Frame(self.root, bg="#2c3e50", height=80)
        header.pack(fill=tk.X)

        title_label = tk.Label(
            header,
            text="🔍 Audit Processor v2.3 (УЛУЧШЕННАЯ ВЕРСИЯ)",
            font=("Arial", 20, "bold"),
            bg="#2c3e50",
            fg="white"
        )
        title_label.pack(pady=20)

        # Основной контейнер
        main_container = tk.Frame(self.root, bg="#f5f5f5")
        main_container.pack(fill=tk.BOTH, expand=True, padx=20, pady=20)

        # Статус AI
        status_frame = tk.Frame(main_container, bg="white", relief=tk.RAISED, borderwidth=1)
        status_frame.pack(fill=tk.X, pady=(0, 15))

        status_color = "#27ae60" if self.ai_available else "#e74c3c"
        status_text = f"✅ Google Gemini ({self.gemini_model})" if self.ai_available else "❌ Gemini не подключен"

        status_label = tk.Label(
            status_frame,
            text=status_text,
            font=("Arial", 12, "bold"),
            bg="white",
            fg=status_color,
            pady=10
        )
        status_label.pack()

        # Секция выбора файлов
        files_frame = tk.LabelFrame(
            main_container,
            text="📁 Выбор файлов",
            font=("Arial", 12, "bold"),
            bg="white",
            padx=15,
            pady=15
        )
        files_frame.pack(fill=tk.X, pady=(0, 15))

        btn_frame = tk.Frame(files_frame, bg="white")
        btn_frame.pack(fill=tk.X, pady=5)

        ttk.Button(
            btn_frame,
            text="📁 Выбрать файлы (изображения/PDF/DOCX)",
            command=self.select_images,
            width=40
        ).pack(side=tk.LEFT, padx=5)

        self.files_listbox = tk.Listbox(
            files_frame,
            height=5,
            font=("Arial", 10),
            bg="#f9f9f9"
        )
        self.files_listbox.pack(fill=tk.BOTH, expand=True, pady=(10, 0))

        # Секция Excel
        excel_frame = tk.LabelFrame(
            main_container,
            text="📊 Шаблон Excel",
            font=("Arial", 12, "bold"),
            bg="white",
            padx=15,
            pady=15
        )
        excel_frame.pack(fill=tk.X, pady=(0, 15))

        excel_btn_frame = tk.Frame(excel_frame, bg="white")
        excel_btn_frame.pack(fill=tk.X)

        self.excel_path_var = tk.StringVar(value="Не выбран")

        ttk.Button(
            excel_btn_frame,
            text="📁 Выбрать Excel",
            command=self.select_excel,
            width=30
        ).pack(side=tk.LEFT, padx=5)

        excel_label = tk.Label(
            excel_btn_frame,
            textvariable=self.excel_path_var,
            font=("Arial", 10),
            bg="white",
            fg="#555"
        )
        excel_label.pack(side=tk.LEFT, padx=10)

        # Кнопка обработки
        self.process_btn = tk.Button(
            main_container,
            text="🚀 НАЧАТЬ ОБРАБОТКУ",
            font=("Arial", 14, "bold"),
            bg="#27ae60",
            fg="white",
            activebackground="#229954",
            activeforeground="white",
            command=self.start_processing,
            height=2,
            cursor="hand2"
        )
        self.process_btn.pack(fill=tk.X, pady=(0, 15))

        # Лог
        log_frame = tk.LabelFrame(
            main_container,
            text="📋 Лог обработки",
            font=("Arial", 12, "bold"),
            bg="white",
            padx=10,
            pady=10
        )
        log_frame.pack(fill=tk.BOTH, expand=True)

        self.log_text = scrolledtext.ScrolledText(
            log_frame,
            height=15,
            font=("Consolas", 9),
            bg="#1e1e1e",
            fg="#00ff00",
            insertbackground="white"
        )
        self.log_text.pack(fill=tk.BOTH, expand=True, pady=(0, 10))

        log_buttons_frame = tk.Frame(log_frame, bg="white")
        log_buttons_frame.pack(fill=tk.X)

        ttk.Button(
            log_buttons_frame,
            text="📋 Копировать логи",
            command=self.copy_logs,
            width=20
        ).pack(side=tk.LEFT, padx=5)

        self.open_file_btn = ttk.Button(
            log_buttons_frame,
            text="📂 Открыть файл",
            command=self.open_result_file,
            width=25,
            state=tk.DISABLED
        )
        self.open_file_btn.pack(side=tk.LEFT, padx=5)

        # Приветствие
        self.log("=" * 70)
        self.log("🔍 Audit Processor УЛУЧШЕННАЯ ВЕРСИЯ v2.3")
        self.log("=" * 70)
        self.log("УЛУЧШЕНИЯ:")
        self.log("  ✅ Полное распознавание текста (не обрывается)")
        self.log("  ✅ Умное сопоставление вопросов-ответов")
        self.log("  ✅ Постобработка OCR для исправления ошибок")
        self.log("  ✅ Улучшенные промпты для AI")
        self.log("  ✅ Полный контекст для AI (500 символов на вопрос, 100 вопросов)")
        self.log("  ✅ Фильтрация OCR-артефактов (лишние символы)")
        self.log("  ✅ Умная классификация документов (взвешенная система)")
        self.log("  ✅ Столбец D с именем файла-источника")
        self.log("  ✅ Поддержка DOC/DOCX/PDF (не только изображения)")
        self.log("")

        if self.ai_available:
            self.log(f"✅ Google Gemini готов ({self.gemini_model})")
        else:
            self.log("❌ ВНИМАНИЕ: Google Gemini не подключен!")
            self.log("💡 Проверьте config.json и API ключ")

        self.log("")

        # Хранение данных
        self.selected_files = []
        self.excel_file = None
        self.excel_header_row = 1
        self.last_created_file = None
        self.is_processing = False

    def log(self, message):
        """Добавить сообщение в лог"""
        self.log_text.insert(tk.END, message + "\n")
        self.log_text.see(tk.END)
        self.log_text.update()

    def copy_logs(self):
        """Копировать логи"""
        logs = self.log_text.get("1.0", tk.END)
        self.root.clipboard_clear()
        self.root.clipboard_append(logs)
        self.root.update()
        messagebox.showinfo("Успех", "✅ Логи скопированы!")

    def open_result_file(self):
        """Открыть готовый файл"""
        if not self.last_created_file or not os.path.exists(self.last_created_file):
            messagebox.showerror("Ошибка", "Файл не найден!")
            return

        try:
            if platform.system() == 'Windows':
                os.startfile(self.last_created_file)
            elif platform.system() == 'Darwin':
                subprocess.run(['open', self.last_created_file])
            else:
                subprocess.run(['xdg-open', self.last_created_file])

            self.log(f"📂 Открыт: {os.path.basename(self.last_created_file)}")
        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось открыть:\n{e}")

    def select_images(self):
        """Выбор файлов (изображения или документы)"""
        files = filedialog.askopenfilenames(
            title="Выберите файлы (изображения или документы)",
            filetypes=[
                ("Все поддерживаемые", "*.jpg *.jpeg *.png *.bmp *.pdf *.docx *.doc"),
                ("Изображения", "*.jpg *.jpeg *.png *.bmp"),
                ("PDF документы", "*.pdf"),
                ("Word документы", "*.docx *.doc"),
                ("Все файлы", "*.*")
            ]
        )

        if files:
            for file in files:
                if file not in self.selected_files:
                    self.selected_files.append(file)
                    # Определяем иконку по типу файла
                    ext = os.path.splitext(file)[1].lower()
                    if ext in ['.pdf']:
                        icon = "📄"
                    elif ext in ['.doc', '.docx']:
                        icon = "📝"
                    else:
                        icon = "🖼️"
                    self.files_listbox.insert(tk.END, f"{icon} {os.path.basename(file)}")

            self.log(f"✅ Добавлено: {len(files)} файлов")

    def select_excel(self):
        """Выбор Excel"""
        file = filedialog.askopenfilename(
            title="Выберите Excel",
            defaultextension=".xlsx",
            filetypes=[("Excel", "*.xlsx *.xls"), ("Все файлы", "*.*")]
        )

        if file:
            try:
                wb = load_workbook(file)
                ws = wb.active

                self.log(f"📊 Анализ: {os.path.basename(file)}")

                header_row = None
                headers = []

                for row_idx in range(1, min(11, ws.max_row + 1)):
                    non_empty = sum(1 for col_idx in range(1, min(ws.max_column + 1, 21))
                                   if ws.cell(row=row_idx, column=col_idx).value)

                    if non_empty >= 2 and not header_row:
                        header_row = row_idx
                        headers = [str(cell.value).strip() for cell in ws[row_idx]
                                  if cell.value and str(cell.value).strip()]

                if not headers:
                    messagebox.showerror("Ошибка", "Заголовки не найдены!")
                    return

                self.excel_file = file
                self.excel_header_row = header_row
                self.excel_path_var.set(os.path.basename(file))

                self.log(f"✅ Выбран: {os.path.basename(file)}")
                self.log(f"   Заголовки в строке: {header_row}")
                self.log(f"   Колонок: {len(headers)}")

            except Exception as e:
                messagebox.showerror("Ошибка", f"Не удалось открыть:\n{e}")

    def query_ai(self, prompt, context=""):
        """Запрос к Google Gemini"""
        full_prompt = f"{context}\n\n{prompt}" if context else prompt

        try:
            response = self.gemini_client.generate_content(
                full_prompt,
                generation_config={
                    "temperature": 0.1,
                    "max_output_tokens": 8192,  # Увеличен лимит
                },
                safety_settings=[
                    {"category": cat, "threshold": "BLOCK_NONE"}
                    for cat in ["HARM_CATEGORY_HARASSMENT", "HARM_CATEGORY_HATE_SPEECH",
                               "HARM_CATEGORY_SEXUALLY_EXPLICIT", "HARM_CATEGORY_DANGEROUS_CONTENT"]
                ]
            )

            # ДЕТАЛЬНОЕ ЛОГИРОВАНИЕ для диагностики
            if response.candidates and len(response.candidates) > 0:
                candidate = response.candidates[0]

                # Проверяем finish_reason
                if hasattr(candidate, 'finish_reason'):
                    finish_reason = str(candidate.finish_reason)
                    if 'SAFETY' in finish_reason:
                        self.log(f"   ⚠️ Ответ заблокирован safety filters: {finish_reason}")
                        # Пытаемся получить safety_ratings
                        if hasattr(candidate, 'safety_ratings'):
                            self.log(f"   Safety ratings: {candidate.safety_ratings}")
                        return "Ошибка: Ответ заблокирован safety filters Gemini"

                if candidate.content and candidate.content.parts:
                    return candidate.content.parts[0].text

                # Если нет parts, пытаемся получить через response.text
                try:
                    if response.text:
                        return response.text
                except Exception as e:
                    self.log(f"   ⚠️ Не удалось получить response.text: {e}")
                    # Логируем детали candidate
                    self.log(f"   Candidate finish_reason: {getattr(candidate, 'finish_reason', 'нет')}")
                    self.log(f"   Candidate content: {getattr(candidate, 'content', 'нет')}")
                    return "Ошибка: пустой ответ Gemini (нет parts)"

            # Если нет candidates
            self.log(f"   ⚠️ Нет candidates в ответе")
            if hasattr(response, 'prompt_feedback'):
                self.log(f"   Prompt feedback: {response.prompt_feedback}")
            return "Ошибка: нет кандидатов ответа (возможно, промпт заблокирован)"

        except Exception as e:
            error_msg = str(e)
            # Проверяем на ошибки лимитов API
            if "quota" in error_msg.lower() or "limit" in error_msg.lower() or "429" in error_msg:
                return "Ошибка: Исчерпан лимит API Gemini. Проверьте квоту на https://aistudio.google.com"
            self.log(f"   ❌ Исключение в query_ai: {error_msg}")
            return f"Ошибка Gemini: {error_msg}"

    def extract_text_from_image(self, file_path):
        """Извлечение текста из изображения через Gemini Vision API"""
        try:
            from PIL import Image

            self.log("   🔍 Gemini Vision OCR...")

            img = Image.open(file_path)
            self.log(f"   📷 Размер: {img.size[0]}x{img.size[1]}px")

            # УЛУЧШЕННЫЙ ПРОМПТ для полного распознавания
            prompt = """Ты - эксперт OCR. Твоя задача - извлечь ВЕСЬ текст с изображения БЕЗ ИСКЛЮЧЕНИЙ.

КРИТИЧЕСКИ ВАЖНО:
1. Распознай КАЖДОЕ слово, КАЖДУЮ букву
2. НЕ ПРОПУСКАЙ ни одной строки
3. Сохрани ВСЕ заголовки, параграфы, списки
4. Читай ДО САМОГО КОНЦА документа
5. Если видишь таблицу - распознай все ячейки
6. Если текст на нескольких страницах - распознай ВСЕ страницы

ФОРМАТ ОТВЕТА:
- Только текст, без комментариев
- Сохрани структуру (заголовки, параграфы)
- Исправь очевидные ошибки OCR

НАЧИНАЙ РАСПОЗНАВАНИЕ:"""

            response = self.gemini_client.generate_content(
                [prompt, img],
                generation_config={
                    "temperature": 0.1,
                    "max_output_tokens": 8192  # Максимум для полного текста
                }
            )

            text = None
            if response.candidates and len(response.candidates) > 0:
                candidate = response.candidates[0]
                if candidate.content and candidate.content.parts:
                    text = candidate.content.parts[0].text

            if not text:
                try:
                    text = response.text
                except:
                    pass

            if text and text.strip():
                # Постобработка
                text = self.post_processor.fix_ocr_errors(text)
                self.log(f"   📝 Распознано: {len(text)} символов")
                return text
            else:
                self.log("   ⚠️ Текст не распознан")
                return "(Текст не найден)"

        except Exception as e:
            error_msg = str(e)
            self.log(f"   ❌ Ошибка OCR: {error_msg}")
            # Проверяем на ошибки лимитов API
            if "quota" in error_msg.lower() or "limit" in error_msg.lower() or "429" in error_msg:
                self.log("   ⚠️ Возможно исчерпан лимит API Gemini")
                return "OCR недоступен (лимит API)"
            return f"Ошибка OCR: {error_msg}"

    def extract_text_from_document(self, file_path):
        """Извлечение текста из DOC/DOCX/PDF документов"""
        file_ext = os.path.splitext(file_path)[1].lower()

        try:
            if file_ext == '.pdf':
                # Извлечение из PDF
                self.log("   📄 Извлечение текста из PDF...")
                try:
                    import PyPDF2
                    text = ""
                    with open(file_path, 'rb') as file:
                        pdf_reader = PyPDF2.PdfReader(file)
                        for page_num in range(len(pdf_reader.pages)):
                            page = pdf_reader.pages[page_num]
                            text += page.extract_text() + "\n"

                    if text.strip():
                        text = self.post_processor.fix_ocr_errors(text)
                        self.log(f"   📝 Извлечено: {len(text)} символов из {len(pdf_reader.pages)} страниц")
                        return text
                    else:
                        self.log("   ⚠️ PDF пустой или текст не извлекается")
                        return "(Текст не извлечен из PDF)"
                except ImportError:
                    self.log("   ❌ PyPDF2 не установлен. Установите: pip install PyPDF2")
                    return "Ошибка: PyPDF2 не установлен"

            elif file_ext in ['.docx', '.doc']:
                # Извлечение из DOCX
                self.log("   📄 Извлечение текста из DOCX...")
                try:
                    from docx import Document
                    doc = Document(file_path)
                    text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])

                    if text.strip():
                        text = self.post_processor.fix_ocr_errors(text)
                        self.log(f"   📝 Извлечено: {len(text)} символов из {len(doc.paragraphs)} параграфов")
                        return text
                    else:
                        self.log("   ⚠️ Документ пустой")
                        return "(Документ пустой)"
                except ImportError:
                    self.log("   ❌ python-docx не установлен. Установите: pip install python-docx")
                    return "Ошибка: python-docx не установлен"

            else:
                return f"Неподдерживаемый формат: {file_ext}"

        except Exception as e:
            self.log(f"   ❌ Ошибка извлечения текста: {e}")
            return f"Ошибка: {e}"

    def match_questions(self, text, table_rows, metadata, source_file=""):
        """Улучшенное сопоставление текста с вопросами таблицы с классификацией документов"""

        # КЛАССИФИКАЦИЯ ДОКУМЕНТА
        doc_type = self.post_processor.classify_document_type(text)
        doc_type_names = {
            'certificate': '📜 Документ об обучении',
            'regulation': '📋 Приказ/Положение/Инструкция',
            'schedule': '📅 График',
            'unknown': '❓ Неизвестный тип'
        }
        self.log(f"   🏷️ Тип документа: {doc_type_names.get(doc_type, 'unknown')}")

        # ИЗВЛЕЧЕНИЕ ТОЛЬКО НУЖНОГО ФРАГМЕНТА
        relevant_fragment = self.post_processor.extract_relevant_fragment(text, doc_type, metadata)
        self.log(f"   📝 Извлечен фрагмент: {len(relevant_fragment)} символов")

        self.log("   🧠 AI анализ соответствия...")

        # УМНАЯ ПРЕДВАРИТЕЛЬНАЯ ФИЛЬТРАЦИЯ вопросов по типу документа
        doc_type_keywords = {
            'certificate': ['обучен', 'квалифик', 'программ', 'курс', 'удостовер', 'компетент'],
            'regulation': ['приказ', 'положен', 'инструкц', 'процедур', 'утвержд', 'регламент'],
            'schedule': ['план', 'график', 'мероприят', 'срок', 'ответствен'],
            'unknown': []  # Для неизвестных типов берем все вопросы
        }

        # Формируем ПОЛНОЕ описание строк с предварительной фильтрацией
        questions_list = []
        keywords = doc_type_keywords.get(doc_type, [])

        for row_num, row_data in table_rows.items():
            # Берем только колонку B (вопросы/информация)
            question_text = ""
            for col_name, col_value in row_data.items():
                if "информацию" in col_name.lower() or "вопрос" in col_name.lower():
                    question_text = col_value
                    break

            if not question_text:
                question_text = next(iter(row_data.values()), "")

            # Пропускаем заголовки
            if len(question_text) < 50:
                continue

            # УМНАЯ ФИЛЬТРАЦИЯ: если есть ключевые слова для типа документа
            if keywords:
                # Проверяем, содержит ли вопрос хотя бы одно ключевое слово
                question_lower = question_text.lower()
                if not any(keyword in question_lower for keyword in keywords):
                    continue  # Пропускаем нерелевантные вопросы

            # Уменьшаем длину вопроса до 300 символов (было 500) для экономии токенов
            questions_list.append(f"Строка {row_num}: {question_text[:300]}")

        # Ограничиваем количество вопросов до 50 (было 100) для избежания MAX_TOKENS
        questions_text = "\n".join(questions_list[:50])

        self.log(f"   📊 Отфильтровано вопросов: {len(questions_list[:50])} из {len(table_rows)}")

        # УЛУЧШЕННЫЙ ПРОМПТ С КЛАССИФИКАЦИЕЙ
        prompt = f"""Ты - эксперт по анализу документов аудита.

ЗАДАЧА: Найди в СПИСКЕ ВОПРОСОВ те, на которые отвечает ДАННЫЙ ФРАГМЕНТ ДОКУМЕНТА.

ТИП ДОКУМЕНТА: {doc_type_names.get(doc_type, 'unknown')}

ВАЖНО! Вставляй в колонку C ТОЛЬКО ПОДХОДЯЩИЙ ФРАГМЕНТ:

ПРАВИЛА ИЗВЛЕЧЕНИЯ ПО ТИПАМ:
- 📋 Приказ/Положение/Инструкция → название, шифр, дата утверждения
- 📅 План/График → название, дата утверждения, название столбцов, один пример мероприятия
- 📜 Документ об обучении → тема обучения, даты обучения, № удостоверения (БЕЗ ФИО!)

СПИСОК ВОПРОСОВ ИЗ ТАБЛИЦЫ:
{questions_text}

ИЗВЛЕЧЕННЫЙ ФРАГМЕНТ ДОКУМЕНТА:
{relevant_fragment}

ПОЛНЫЙ ТЕКСТ (для контекста):
{text[:800]}

МЕТАДАННЫЕ:
- Организация: {metadata.get('organization', 'не указано')}
- Тип: {metadata.get('doc_type', 'не указано')}
- Номер: {metadata.get('doc_number', 'не указано')}
- Дата: {metadata.get('doc_date', 'не указано')}

ИНСТРУКЦИЯ:
1. Внимательно прочитай ФРАГМЕНТ ДОКУМЕНТА
2. Найди в СПИСКЕ ВОПРОСОВ те, на которые этот фрагмент дает ответ
3. Верни НОМЕРА СТРОК (от 1 до 3 наиболее подходящих)
4. В колонку C будет вставлен ТОЛЬКО ФРАГМЕНТ, а не весь текст

ФОРМАТ ОТВЕТА (строго JSON, БЕЗ лишнего текста):
{{"matched_rows": [123, 145], "confidence": "высокая", "reason": "фрагмент содержит...", "fragment_to_insert": "краткое извлечение"}}

JSON:"""

        try:
            response = self.query_ai(prompt)

            # ОТЛАДКА: Показываем ответ AI
            self.log(f"   🤖 Ответ AI (первые 300 символов): {response[:300]}")

            # Извлекаем JSON
            json_match = re.search(r'\{[\s\S]*?"matched_rows"[\s\S]*?\}', response)

            if json_match:
                self.log(f"   ✓ JSON найден в ответе")
                try:
                    result = json.loads(json_match.group(0))

                    if "matched_rows" in result:
                        rows = [int(r) for r in result['matched_rows'] if isinstance(r, (int, str)) and str(r).isdigit()]

                        self.log(f"   ✓ Найдено соответствий: {len(rows)}")
                        self.log(f"     Строки: {rows}")
                        self.log(f"     Уверенность: {result.get('confidence', 'не указана')}")

                        return {
                            "matched_rows": rows,
                            "confidence": result.get('confidence', 'средняя'),
                            "reason": result.get('reason', 'AI определил соответствие'),
                            "fragment": relevant_fragment,  # Используем извлеченный фрагмент
                            "doc_type": doc_type,
                            "source_file": source_file  # Имя файла-источника
                        }

                except json.JSONDecodeError as je:
                    self.log(f"   ⚠️ Ошибка парсинга JSON: {je}")
                    self.log(f"   JSON строка: {json_match.group(0)[:200]}")
            else:
                self.log("   ⚠️ JSON не найден в ответе AI")

            self.log("   ⚠️ AI не нашел соответствий")
            return None

        except Exception as e:
            self.log(f"   ❌ Ошибка: {e}")
            import traceback
            self.log(f"   Traceback: {traceback.format_exc()[:300]}")
            return None

    def start_processing(self):
        """Начать обработку"""

        if self.is_processing:
            messagebox.showwarning("Внимание", "⚠️ Обработка уже выполняется!")
            return

        if not self.selected_files:
            messagebox.showwarning("Внимание", "Выберите файлы!")
            return

        if not self.excel_file:
            messagebox.showwarning("Внимание", "Выберите Excel!")
            return

        if not self.ai_available:
            result = messagebox.askyesno(
                "AI недоступен",
                "AI не подключен. Продолжить в демо-режиме?"
            )
            if not result:
                return

        self.is_processing = True
        self.process_btn.config(state=tk.DISABLED, text="⏳ ОБРАБОТКА...", bg="#95a5a6")

        thread = threading.Thread(target=self.process_files, daemon=True)
        thread.start()

    def read_table_rows(self, ws, header_row_num, headers):
        """Читает строки таблицы"""
        table_rows = {}

        header_positions = {}
        for idx, header in enumerate(headers, start=1):
            header_positions[header] = idx

        for row_idx in range(header_row_num + 1, ws.max_row + 1):
            row_data = {}
            has_content = False

            for col_name, col_idx in header_positions.items():
                cell_value = ws.cell(row=row_idx, column=col_idx).value
                if cell_value and str(cell_value).strip():
                    row_data[col_name] = str(cell_value).strip()
                    has_content = True

            if has_content:
                table_rows[row_idx] = row_data

        return table_rows, header_positions

    def process_files(self):
        """Основная обработка файлов"""

        import time
        start_time = time.time()

        self.log("\n" + "=" * 70)
        self.log("🚀 НАЧАЛО ОБРАБОТКИ (УЛУЧШЕННАЯ ВЕРСИЯ)")
        self.log("=" * 70)

        # Создание выходного файла
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        base_dir = os.path.dirname(self.excel_file)
        base_name = os.path.splitext(os.path.basename(self.excel_file))[0]
        output_file = os.path.join(base_dir, f"{base_name}_заполнен_{timestamp}.xlsx")

        try:
            wb = load_workbook(self.excel_file)
            ws = wb.active

            # Читаем заголовки
            headers = []
            for cell in ws[self.excel_header_row]:
                if cell.value and str(cell.value).strip():
                    headers.append(str(cell.value).strip())

            self.log(f"📊 Колонок: {len(headers)}")

            # Читаем строки таблицы
            table_rows, header_positions = self.read_table_rows(ws, self.excel_header_row, headers)
            self.log(f"📋 Строк с данными: {len(table_rows)}")

            # Добавляем колонки для свидетельств и источника если нет
            if "Свидетельства" not in header_positions:
                headers.append("Свидетельства")
                col_c = 3
                ws.cell(row=self.excel_header_row, column=col_c, value="Свидетельства")
                header_positions["Свидетельства"] = col_c
                self.log("   Добавлена колонка 'Свидетельства' (C)")

            if "Источник" not in header_positions:
                headers.append("Источник")
                col_d = 4
                ws.cell(row=self.excel_header_row, column=col_d, value="Источник")
                header_positions["Источник"] = col_d
                self.log("   Добавлена колонка 'Источник' (D)")

            # Счетчики
            matched_count = 0
            updated_rows = []

            # Обработка файлов
            for idx, file_path in enumerate(self.selected_files, start=1):
                file_name = os.path.basename(file_path)
                file_ext = os.path.splitext(file_path)[1].lower()
                self.log(f"\n📄 [{idx}/{len(self.selected_files)}] {file_name}")

                # Определяем тип файла и извлекаем текст
                text = None
                if file_ext in ['.jpg', '.jpeg', '.png', '.bmp']:
                    # Изображение - используем OCR
                    text = self.extract_text_from_image(file_path)
                elif file_ext in ['.pdf', '.doc', '.docx']:
                    # Документ - извлекаем текст напрямую
                    text = self.extract_text_from_document(file_path)
                else:
                    self.log(f"   ⚠️ Неподдерживаемый формат: {file_ext}")
                    continue

                if not text or len(text.strip()) < 10:
                    self.log("   ⚠️ Мало текста, пропускаем")
                    continue

                # Извлекаем метаданные
                metadata = self.post_processor.extract_metadata(text)

                if metadata.get('organization'):
                    self.log(f"   Организация: {metadata['organization']}")
                if metadata.get('doc_type'):
                    self.log(f"   Тип: {metadata['doc_type']}")

                # Сопоставление
                if self.ai_available and table_rows:
                    match_result = self.match_questions(text, table_rows, metadata, source_file=file_name)

                    if match_result and match_result.get("matched_rows"):
                        rows = match_result["matched_rows"]
                        reason = match_result.get("reason", "")
                        fragment = match_result.get("fragment", text[:300])  # Используем фрагмент
                        doc_type = match_result.get("doc_type", "unknown")

                        for row_num in rows:
                            if row_num in table_rows:
                                # Вставляем в колонку Свидетельства ТОЛЬКО ФРАГМЕНТ
                                col_c_idx = header_positions.get("Свидетельства", 3)
                                col_d_idx = header_positions.get("Источник", 4)

                                # Колонка C - фрагмент
                                existing_fragment = ws.cell(row=row_num, column=col_c_idx).value
                                new_fragment = f"{existing_fragment}\n\n{fragment}" if existing_fragment else fragment
                                ws.cell(row=row_num, column=col_c_idx, value=new_fragment)

                                # Колонка D - имя файла (источник)
                                existing_source = ws.cell(row=row_num, column=col_d_idx).value
                                new_source = f"{existing_source}\n{file_name}" if existing_source else file_name
                                ws.cell(row=row_num, column=col_d_idx, value=new_source)

                                self.log(f"   ✓ Добавлено в строку {row_num}")
                                updated_rows.append(row_num)

                        matched_count += 1
                        self.log(f"✅ Размещено в {len(rows)} строках")
                    else:
                        self.log("⚠️ Соответствий не найдено")

            # Сохранение
            self.log(f"\n💾 Сохранение: {os.path.basename(output_file)}")
            wb.save(output_file)
            self.last_created_file = output_file

            elapsed = time.time() - start_time
            minutes = int(elapsed // 60)
            seconds = int(elapsed % 60)

            self.log("\n" + "=" * 70)
            self.log("✅ ОБРАБОТКА ЗАВЕРШЕНА!")
            self.log("=" * 70)
            self.log(f"⏱️  Время: {minutes} мин {seconds} сек" if minutes > 0 else f"⏱️  Время: {seconds} сек")
            self.log(f"📂 Файл: {output_file}")
            self.log(f"📝 Обработано: {len(self.selected_files)} файлов")
            self.log(f"✓ Найдено соответствий: {matched_count}")
            self.log(f"✓ Обновлено строк: {len(set(updated_rows))}")

            self.open_file_btn.config(state=tk.NORMAL)

            messagebox.showinfo(
                "Успех",
                f"✅ Готово!\n\nФайлов: {len(self.selected_files)}\nСоответствий: {matched_count}\nОбновлено строк: {len(set(updated_rows))}\n\nНажмите '📂 Открыть файл'"
            )

        except Exception as e:
            self.log(f"\n❌ ОШИБКА: {e}")
            import traceback
            self.log(f"{traceback.format_exc()[:500]}")
            messagebox.showerror("Ошибка", f"Ошибка:\n{e}")

        finally:
            self.is_processing = False
            self.process_btn.config(state=tk.NORMAL, text="🚀 НАЧАТЬ ОБРАБОТКУ", bg="#27ae60")


def main():
    """Точка входа"""

    print("=" * 70)
    print("🔍 Audit Processor УЛУЧШЕННАЯ ВЕРСИЯ v2.3")
    print("=" * 70)
    print()
    print("УЛУЧШЕНИЯ:")
    print("  ✅ Полное распознавание текста (не обрывается)")
    print("  ✅ Умное сопоставление вопросов-ответов")
    print("  ✅ Постобработка OCR для исправления ошибок")
    print("  ✅ Улучшенные промпты для AI")
    print("  ✅ Полный контекст для AI (500 символов на вопрос, до 100 вопросов)")
    print("  ✅ Фильтрация OCR-артефактов (лишние символы)")
    print("  ✅ Умная классификация документов (взвешенная система)")
    print("  ✅ Столбец D с именем файла-источника")
    print("  ✅ Поддержка DOC/DOCX/PDF (не только изображения)")
    print()
    print("=" * 70)
    print()

    root = tk.Tk()
    app = AuditProcessorApp(root)
    root.mainloop()


if __name__ == "__main__":
    main()
